from docx import Document
from docx.enum.text import WD_LINE_SPACING
from datetime import date
from sys import argv
import csv

def main():
    if len(argv) != 2:
        print("Usage: python application.py <csv file path>")
        return 1

    today = date.today().strftime("%B %d, %Y")
    doc_names = []

    with open(argv[1], "r") as file:
        reader = csv.DictReader(file)
        i = 0
        for row in reader:
            document = Document()
            names = row['name'].split()
            address = row['address'].split(", ", maxsplit=1)

            doc_names.append(f"thank-you-to-{'-'.join(names)}")

            document.add_paragraph("\n\n\n{}\n\n\n".format(today))
            document.add_paragraph(row['name'].upper())
            document.add_paragraph(address[0].upper())
            document.add_paragraph(address[1].upper())
            document.add_paragraph(f"\nDear {names[0]},")
            document.add_paragraph("Thank you for attending the wedding. Jane and I both really appreciate the gift you got us.")
            document.save(f'{doc_names[i]}.docx')
            i+=1

main()